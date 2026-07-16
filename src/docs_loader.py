"""
data/docs/ altindaki kullanici dosyalarini RAG chunk'larina cevirir.

Desteklenen:
  - .docx  (Word: paragraf + tablo)
  - .txt / .md
  - .csv   (satir satirlari metin chunk)

Kullanim:
  docs'u data/docs/ icine at
  python src/rag_pipeline.py build
"""
from __future__ import annotations

import csv
import re
from pathlib import Path
from typing import Optional

# BIST sembolleri (dosya adi / metin icinden yakalamak icin)
SYMBOL_RE = re.compile(
    r"\b(THYAO|ASELS|GARAN|AKBNK|ISCTR|YKBNK|TUPRS|EREGL|KCHOL|SAHOL|"
    r"TOASO|FROTO|ARCLK|PGSUS|TCELL|TTKOM|BIMAS|MGROS|ULKER|SISE|"
    r"SISE|PETKM|KOZAL|ENKAI|SASA|HEKTS|VESTL|DOAS|GUBRF)\b",
    re.IGNORECASE,
)


def _detect_symbol(text: str, filename: str = "") -> Optional[str]:
    for src in (filename, text[:2000]):
        m = SYMBOL_RE.search(src or "")
        if m:
            return m.group(1).upper()
    return None


def _chunk_text(
    text: str,
    chunk_size: int = 800,
    chunk_overlap: int = 100,
) -> list[str]:
    """Paragraf sinirlarina saygi duyan basit chunker."""
    text = re.sub(r"\n{3,}", "\n\n", (text or "").strip())
    if not text:
        return []
    if len(text) <= chunk_size:
        return [text]

    paras = [p.strip() for p in text.split("\n\n") if p.strip()]
    chunks: list[str] = []
    buf = ""
    for p in paras:
        if not buf:
            buf = p
        elif len(buf) + 2 + len(p) <= chunk_size:
            buf = f"{buf}\n\n{p}"
        else:
            chunks.append(buf)
            # overlap: onceki chunk'in sonu
            if chunk_overlap > 0 and len(buf) > chunk_overlap:
                tail = buf[-chunk_overlap:]
                buf = f"{tail}\n\n{p}"
            else:
                buf = p
            # tek paragraf cok uzunsa zorla bol
            while len(buf) > chunk_size * 2:
                chunks.append(buf[:chunk_size])
                buf = buf[chunk_size - chunk_overlap :]
    if buf:
        chunks.append(buf)
    return chunks


def extract_docx(path: Path) -> str:
    """Word dosyasindan duz metin + tablolari cek."""
    try:
        from docx import Document
    except ImportError as e:
        raise ImportError(
            "python-docx yok. Kur: pip install python-docx"
        ) from e

    doc = Document(str(path))
    parts: list[str] = []

    for p in doc.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)

    for ti, table in enumerate(doc.tables):
        rows_out = []
        for row in table.rows:
            cells = [" ".join((c.text or "").split()) for c in row.cells]
            # birlesik hucre tekrari temizle (Word merge)
            cleaned = []
            prev = None
            for c in cells:
                if c != prev:
                    cleaned.append(c)
                    prev = c
            if any(cleaned):
                rows_out.append(" | ".join(cleaned))
        if rows_out:
            parts.append(f"[Tablo {ti + 1}]\n" + "\n".join(rows_out))

    return "\n\n".join(parts)


def extract_text_file(path: Path) -> str:
    return path.read_text(encoding="utf-8", errors="replace")


def extract_csv(path: Path) -> str:
    lines = []
    with open(path, newline="", encoding="utf-8", errors="replace") as f:
        reader = csv.reader(f)
        for row in reader:
            if any(cell.strip() for cell in row):
                lines.append(" | ".join(cell.strip() for cell in row))
    return "\n".join(lines)


def load_docs_dir(
    docs_dir: str | Path = "data/docs",
    chunk_size: int = 800,
    chunk_overlap: int = 100,
) -> tuple[list[str], list[dict]]:
    """
    data/docs icindeki dosyalari oku -> (texts, metas).
    Bos klasor / dosya yoksa bos liste doner.
    """
    root = Path(docs_dir)
    if not root.exists():
        root.mkdir(parents=True, exist_ok=True)
        return [], []

    texts: list[str] = []
    metas: list[dict] = []

    patterns = ("*.docx", "*.txt", "*.md", "*.csv")
    files: list[Path] = []
    for pat in patterns:
        files.extend(sorted(root.rglob(pat)))
    # temp Word kilit dosyalarini atla
    files = [f for f in files if not f.name.startswith("~$")]

    for path in files:
        try:
            suffix = path.suffix.lower()
            if suffix == ".docx":
                raw = extract_docx(path)
                ftype = "docx"
            elif suffix == ".csv":
                raw = extract_csv(path)
                ftype = "csv"
            else:
                raw = extract_text_file(path)
                ftype = "text"
        except Exception as e:
            print(f"  x docs {path.name}: {e}")
            continue

        if not (raw or "").strip():
            print(f"  ~ docs {path.name}: bos, atlandi")
            continue

        rel = str(path.relative_to(root)) if path.is_relative_to(root) else path.name
        symbol = _detect_symbol(raw, path.stem)
        header = f"Kaynak dosya: {rel}"
        if symbol:
            header += f" | Sembol: {symbol}"

        chunks = _chunk_text(raw, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        for i, ch in enumerate(chunks):
            body = f"{header}\n\n{ch}"
            texts.append(body)
            metas.append(
                {
                    "type": "user_doc",
                    "format": ftype,
                    "source": rel,
                    "chunk": i,
                    "symbol": symbol or "",
                }
            )
        print(f"  + docs {rel}: {len(chunks)} chunk" + (f" [{symbol}]" if symbol else ""))

    return texts, metas
